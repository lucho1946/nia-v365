"""
file_processor.py — Extrae lista de ítems de cualquier archivo enviado por el cliente.
Soporta: Excel (.xlsx/.xls/.csv), PDF, Word (.docx), imágenes (.jpg/.png/.webp)
Cada ítem retornado tiene: texto, fila (opcional), campo_detectado
"""

import io
import os
import base64
from typing import List, Dict
import httpx

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# ─── Dispatcher principal ─────────────────────────────────────────────────────
async def procesar_archivo(contenido: bytes, nombre: str) -> List[Dict]:
    """
    Detecta tipo de archivo por extensión y llama al extractor correcto.
    Retorna lista de dicts: [{texto, fila, campo_detectado}]
    """
    ext = nombre.lower().split(".")[-1]

    if ext in ("xlsx", "xls"):
        return extraer_excel(contenido)
    elif ext == "csv":
        return extraer_csv(contenido)
    elif ext == "pdf":
        return extraer_pdf(contenido)
    elif ext in ("docx", "doc"):
        return extraer_word(contenido)
    elif ext in ("jpg", "jpeg", "png", "webp"):
        return await extraer_imagen(contenido, ext)
    else:
        # Intenta como texto plano
        return extraer_texto_plano(contenido)

# ─── Excel ────────────────────────────────────────────────────────────────────
def extraer_excel(contenido: bytes) -> List[Dict]:
    import openpyxl
    wb  = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
    ws  = wb.active
    items = []

    # Detecta encabezados en fila 1
    headers = []
    for cell in ws[1]:
        headers.append(str(cell.value).lower().strip() if cell.value else "")

    CAMPOS_CODIGO = {"codigo", "código", "code", "ref", "referencia", "sku", "part", "part number"}
    CAMPOS_NOMBRE = {"nombre", "name", "descripcion", "descripción", "producto", "item", "description"}
    CAMPOS_CANT   = {"cantidad", "qty", "quantity", "cant"}

    col_codigo = next((i for i, h in enumerate(headers) if h in CAMPOS_CODIGO), None)
    col_nombre = next((i for i, h in enumerate(headers) if h in CAMPOS_NOMBRE), None)
    col_cant   = next((i for i, h in enumerate(headers) if h in CAMPOS_CANT), None)

    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not any(row):
            continue
        partes = []
        campo  = "texto_libre"

        if col_codigo is not None and row[col_codigo]:
            partes.append(str(row[col_codigo]).strip())
            campo = "codigo"
        if col_nombre is not None and row[col_nombre]:
            partes.append(str(row[col_nombre]).strip())
            if campo == "texto_libre":
                campo = "nombre"
        if not partes:
            partes = [str(v).strip() for v in row if v]

        cantidad = None
        if col_cant is not None and row[col_cant]:
            try:
                cantidad = int(float(str(row[col_cant])))
            except Exception:
                pass

        texto = " ".join(partes).strip()
        if texto:
            items.append({
                "texto":            texto,
                "fila":             row_idx,
                "campo_detectado":  campo,
                "cantidad":         cantidad,
            })
    return items

# ─── CSV ──────────────────────────────────────────────────────────────────────
def extraer_csv(contenido: bytes) -> List[Dict]:
    import csv
    texto = contenido.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(texto))
    rows   = list(reader)
    if not rows:
        return []
    # Reutiliza lógica de excel convirtiendo a openpyxl-like estructura
    # Simplificado: cada fila como texto libre
    items = []
    for idx, row in enumerate(rows[1:], start=2):
        texto = " ".join(c.strip() for c in row if c.strip())
        if texto:
            items.append({"texto": texto, "fila": idx, "campo_detectado": "texto_libre", "cantidad": None})
    return items

# ─── PDF ──────────────────────────────────────────────────────────────────────
def extraer_pdf(contenido: bytes) -> List[Dict]:
    import fitz  # PyMuPDF
    doc   = fitz.open(stream=contenido, filetype="pdf")
    items = []
    linea_idx = 0

    for page in doc:
        texto_pagina = page.get_text("text")
        for linea in texto_pagina.splitlines():
            linea = linea.strip()
            if len(linea) < 3:
                continue
            linea_idx += 1
            items.append({
                "texto":           linea,
                "fila":            linea_idx,
                "campo_detectado": "texto_libre",
                "cantidad":        None,
            })

    # Si no extrajo texto (PDF escaneado) → retorna ítem especial para OCR
    if not items:
        items = [{"texto": "[PDF escaneado — requiere OCR manual]", "fila": 1,
                  "campo_detectado": "ocr_requerido", "cantidad": None}]
    return items

# ─── Word ─────────────────────────────────────────────────────────────────────
def extraer_word(contenido: bytes) -> List[Dict]:
    from docx import Document
    doc   = Document(io.BytesIO(contenido))
    items = []
    idx   = 0

    # Párrafos
    for para in doc.paragraphs:
        texto = para.text.strip()
        if len(texto) >= 3:
            idx += 1
            items.append({"texto": texto, "fila": idx, "campo_detectado": "texto_libre", "cantidad": None})

    # Tablas — cada fila como ítem
    for tabla in doc.tables:
        headers = [c.text.strip().lower() for c in tabla.rows[0].cells] if tabla.rows else []
        for row in tabla.rows[1:]:
            celdas = [c.text.strip() for c in row.cells if c.text.strip()]
            if celdas:
                idx += 1
                items.append({
                    "texto":           " ".join(celdas),
                    "fila":            idx,
                    "campo_detectado": "tabla",
                    "cantidad":        None,
                })
    return items

# ─── Imagen — GPT-4o-mini Vision ─────────────────────────────────────────────
async def extraer_imagen(contenido: bytes, ext: str) -> List[Dict]:
    """
    Envía la imagen a GPT-4o-mini Vision.
    El modelo extrae códigos, referencias y nombres visibles.
    """
    media_type = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
    b64        = base64.b64encode(contenido).decode("utf-8")

    payload = {
        "model": "gpt-4o-mini",
        "max_tokens": 800,
        "messages": [{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Eres un asistente de compras industriales. "
                        "Extrae todos los productos, códigos, referencias o nombres de equipos visibles en esta imagen. "
                        "Devuelve SOLO una lista, un ítem por línea, sin explicaciones. "
                        "Si hay cantidades, ponlas al final del ítem así: 'nombre del producto | cantidad: N'. "
                        "Si no hay productos visibles, responde: SIN_PRODUCTOS."
                    )
                },
                {
                    "type":      "image_url",
                    "image_url": {"url": f"data:{media_type};base64,{b64}"}
                }
            ]
        }]
    }

    try:
        async with httpx.AsyncClient(timeout=20) as client:
            r = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json=payload,
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}"}
            )
            texto = r.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        return [{"texto": "[Error procesando imagen]", "fila": 1,
                 "campo_detectado": "error", "cantidad": None}]

    if texto == "SIN_PRODUCTOS":
        return [{"texto": "[Imagen sin productos identificables]", "fila": 1,
                 "campo_detectado": "sin_resultado", "cantidad": None}]

    items = []
    for idx, linea in enumerate(texto.splitlines(), start=1):
        linea = linea.strip().lstrip("-•*").strip()
        if not linea:
            continue
        cantidad = None
        if "| cantidad:" in linea.lower():
            partes = linea.lower().split("| cantidad:")
            linea  = partes[0].strip()
            try:
                cantidad = int(partes[1].strip())
            except Exception:
                pass
        items.append({"texto": linea, "fila": idx, "campo_detectado": "vision", "cantidad": cantidad})

    return items

# ─── Texto plano ──────────────────────────────────────────────────────────────
def extraer_texto_plano(contenido: bytes) -> List[Dict]:
    texto = contenido.decode("utf-8", errors="ignore")
    items = []
    for idx, linea in enumerate(texto.splitlines(), start=1):
        linea = linea.strip()
        if len(linea) >= 3:
            items.append({"texto": linea, "fila": idx, "campo_detectado": "texto_libre", "cantidad": None})
    return items
